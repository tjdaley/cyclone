"""
app/services/docx_service.py - Generate Word documents for discovery responses.

Produces a .docx file with one section per discovery_request_item,
formatted for paralegal review and filing. Parses basic markdown
formatting (bold, italic, numbered/bulleted lists) into native Word runs.
"""
import io
import re

from util.loggerfactory import LoggerFactory

LOGGER = LoggerFactory.create_logger(__name__)


def _slug_to_title(slug: str) -> str:
    """Convert a slug like 'attorney-client' to 'Attorney Client'."""
    return slug.replace("-", " ").title()


_NUMBERED_ITEM = re.compile(r"^\s*(\d+)\.\s+(.*)")
_BULLET_ITEM = re.compile(r"^\s*[-*]\s+(.*)")

_BLOCK_SPACING_PT = 10  # Vertical space between paragraphs (none within one)


def _is_list_item(line: str) -> bool:
    """
    :param line: A raw line of markdown.
    :type line: str
    :return: True if the line opens a numbered or bulleted list item.
    :rtype: bool
    """
    return bool(_NUMBERED_ITEM.match(line) or _BULLET_ITEM.match(line))


def _add_markdown_text(doc, text: str) -> None:
    """
    Parse markdown text and add it to the document as formatted paragraphs.

    Line handling follows standard markdown block rules, which is what lets a
    witness block stay tight while blocks stay separated:

    - A single newline is a soft break *inside* the paragraph — the lines are
      single-spaced against each other (Word ``<w:br/>``, not a new ``<w:p>``).
    - A blank line ends the paragraph; the next one starts after
      ``_BLOCK_SPACING_PT`` of space.

    Also supports **bold**/*italic* and numbered/bulleted lists. A list item
    always begins its own paragraph, so a list directly under a line of text
    is not swallowed into that paragraph.
    """
    from docx.enum.text import WD_BREAK  # noqa: PLC0415
    from docx.shared import Pt  # noqa: PLC0415

    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Blank lines separate paragraphs; the spacing comes from space_after.
        if not stripped:
            i += 1
            continue

        num_match = _NUMBERED_ITEM.match(line)
        if num_match:
            p = doc.add_paragraph(style="List Number")
            _add_inline_formatting(p, num_match.group(2))
            i += 1
            continue

        bullet_match = _BULLET_ITEM.match(line)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_formatting(p, bullet_match.group(1))
            i += 1
            continue

        # Gather every following line up to the next blank line or list item —
        # they belong to this same paragraph, separated by soft line breaks.
        block = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not _is_list_item(lines[i]):
            block.append(lines[i].strip())
            i += 1

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(_BLOCK_SPACING_PT)
        for position, part in enumerate(block):
            if position:
                p.add_run().add_break(WD_BREAK.LINE)
            _add_inline_formatting(p, part)


def _add_inline_formatting(paragraph, text: str) -> None:
    """
    Parse inline markdown (**bold**, *italic*, ***bold italic***) and add
    formatted runs to the paragraph.
    """
    # Pattern matches: ***bold italic***, **bold**, *italic*, or plain text
    pattern = re.compile(r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)")

    pos = 0
    for m in pattern.finditer(text):
        # Add plain text before this match
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])

        if m.group(2):  # ***bold italic***
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.italic = True
        elif m.group(3):  # **bold**
            run = paragraph.add_run(m.group(3))
            run.bold = True
        elif m.group(4):  # *italic*
            run = paragraph.add_run(m.group(4))
            run.italic = True

        pos = m.end()

    # Add remaining plain text after last match
    if pos < len(text):
        paragraph.add_run(text[pos:])


def generate_discovery_response_docx(
    document_type: str,
    matter_name: str,
    items: list[dict],
) -> bytes:
    """
    Generate a Word document containing formatted discovery responses.

    :param document_type: Type of discovery (e.g. 'interrogatories').
    :param matter_name: Matter name for the document title.
    :param items: List of discovery_request_item dicts, each containing
                  request_number, source_text, interpretations, privileges,
                  objections, and response.
    :return: Raw bytes of the .docx file.
    """
    from docx import Document  # noqa: PLC0415
    from docx.shared import Pt  # noqa: PLC0415
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: PLC0415

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Title
    title = doc.add_heading(level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"Responses to {_slug_to_title(document_type)}")
    run.font.size = Pt(14)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(matter_name).italic = True
    doc.add_paragraph()  # spacer

    for item in items:
        request_number = item.get("request_number", "?")
        source_text = item.get("source_text", "")
        interpretations = item.get("interpretations", [])
        privileges = item.get("privileges", [])
        objections = item.get("objections", [])
        response = item.get("response")

        # Request header + text
        p = doc.add_paragraph()
        run = p.add_run(f"Request #{request_number}: ")
        run.bold = True
        _add_inline_formatting(p, source_text.replace("\n", " "))

        # Interpretations
        if interpretations:
            doc.add_paragraph()
            h = doc.add_paragraph()
            h.add_run("Interpretations:").bold = True
            for idx, interp in enumerate(interpretations, 1):
                p = doc.add_paragraph(style="List Number")
                _add_inline_formatting(p, interp)

        # Privileges
        if privileges:
            doc.add_paragraph()
            h = doc.add_paragraph()
            h.add_run("Privileges:").bold = True
            for priv in privileges:
                name = _slug_to_title(priv.get("privilege_name", ""))
                text = priv.get("text", "")
                p = doc.add_paragraph()
                p.add_run(f"{name}: ").bold = True
                _add_inline_formatting(p, text)

        # Objections
        if objections:
            doc.add_paragraph()
            h = doc.add_paragraph()
            h.add_run("Objections:").bold = True
            for obj in objections:
                name = _slug_to_title(obj.get("objection_name", ""))
                text = obj.get("text", "")
                p = doc.add_paragraph()
                p.add_run(f"{name}: ").bold = True
                _add_inline_formatting(p, text)

        # Response
        if response and response.strip():
            doc.add_paragraph()
            h = doc.add_paragraph()
            h.add_run("Response:").bold = True
            _add_markdown_text(doc, response)

        # Blank line between items
        doc.add_paragraph()

    # Write to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    LOGGER.info("docx_service: generated document with %s items", len(items))
    return buf.getvalue()
