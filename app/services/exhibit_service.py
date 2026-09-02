"""
app/services/exhibit_service.py - Turning query results into a court exhibit.

One document description, four renderings. An ``Exhibit`` says what the document
contains — caption, what was selected, the table, the totals — and each renderer
says how that looks in its own format. Anything in Cyclone that produces a table
worth taking to court builds an ``Exhibit``; none of them learn to write DOCX.

The caption is a **template**, not renderer code. Firms disagree about captions,
and the disagreement is always about wording and order, never about how bold
text is written into a .docx. Keeping the template as a handful of format
strings means the future "let this firm edit its own caption" release changes
data rather than three renderers — see ``_SYSTEM_CAPTION`` below.

CSV is deliberately not an exhibit. It is the clean extraction: header row, data
rows, nothing else, so it can be read by a spreadsheet or handed to a model
without a preamble to strip first. The exhibit formats carry the caption and the
verification notice; the CSV carries the data.
"""
import csv
import html
import io
import re
from dataclasses import dataclass, field
from typing import Any, Optional

from db.models.matter import ClientAlignment
from util.loggerfactory import LoggerFactory

LOGGER = LoggerFactory.create_logger(__name__)

# A blank on a court document is a rule, not the word "None". Anything the
# matter cannot supply is printed as a fill-in and reported in `warnings`.
_BLANK = "__________"

# Under Texas Rule of Evidence 1006 a summary may stand in for voluminous
# records only if the originals are available and the summary is accurate — and
# the proponent carries that accuracy, not the tool that drew the table. The
# second sentence is the specific one: extraction is measured, not perfect, and
# a wrong date inside a statement period is exactly the error that reconciles
# cleanly and reaches an exhibit unflagged.
NOTICE = (
    "This exhibit is a summary prepared from records produced in this case. The "
    "underlying records must be made available to the other parties, and the party "
    "offering this summary is responsible for its accuracy. Entries were extracted "
    "from the produced documents by automated means and may contain errors. Verify "
    "every date, amount, and description against the original documents before this "
    "exhibit is shown to anyone or offered in court."
)


# ── The document description ─────────────────────────────────────────────────

@dataclass(frozen=True)
class Run:
    """A stretch of text with one set of styles. The unit every renderer draws."""
    text: str
    bold: bool = False
    underline: bool = False


@dataclass(frozen=True)
class Line:
    """One line of the caption."""
    runs: tuple[Run, ...]
    align: str = "center"
    blank_after: bool = False


@dataclass(frozen=True)
class Column:
    """
    A table column.

    ``numeric`` right-aligns it. ``money`` additionally formats it as currency
    **in the exhibit formats only** — the CSV keeps the raw value, because a
    spreadsheet given "-$2,500.00" reads it as text and will not add it up.
    """
    heading: str
    numeric: bool = False
    money: bool = False


@dataclass
class Exhibit:
    """
    A table with everything needed to put it in front of a court.

    ``selection`` is what makes this usable by someone — or something — that did
    not run the query: a table of forty transactions means nothing without the
    criteria that produced it. It is also the part a model needs most when the
    markdown is pasted in and asked for a print-ready exhibit.
    """
    name: str
    caption: tuple[Line, ...] = ()
    columns: tuple[Column, ...] = ()
    rows: tuple[tuple[str, ...], ...] = ()
    selection: tuple[tuple[str, str], ...] = ()
    summary: tuple[tuple[str, str], ...] = ()
    #: Marks that qualify individual rows — "† institution inferred". They ride
    #: with the table rather than in `selection`, because a reader meeting a
    #: dagger in a cell looks directly below the table for what it means, and a
    #: mark whose explanation is missing is worse than no mark at all.
    footnotes: tuple[str, ...] = ()
    notice: str = NOTICE
    warnings: list[str] = field(default_factory=list)

    @property
    def filename_stem(self) -> str:
        """A safe filename built from the exhibit's own name."""
        stem = re.sub(r"[^A-Za-z0-9]+", "_", self.name).strip("_")
        return stem or "exhibit"


# ── The caption template ─────────────────────────────────────────────────────

# The firm-wide caption. Each entry is (template, blank_after).
#
# `**bold**` and `__underline__` mark styles; `{name}` interpolates from the
# context built by `_caption_context`. A line whose text resolves to nothing is
# dropped, which is how a matter with no case style on file still produces a
# usable heading rather than a stranded blank line.
#
# FUTURE: per-firm and per-user overrides. The override belongs in a table keyed
# the way matter_preferences will be — a NULL user id meaning "the firm's
# template" and a row meaning "this user's" — and it stores exactly this: a list
# of template strings. Nothing below this constant needs to change to allow it;
# `caption_lines` already takes the template as an argument for that reason.
_SYSTEM_CAPTION: tuple[tuple[str, bool], ...] = (
    ("**Cause No: __{cause_number}__**", False),
    ("**In the {court_name} of {county} County, {state}**", False),
    ("{case_style}", True),
    ("**{alignment_possessive}{exhibit_name}**", True),
)

_MARKUP = re.compile(r"\*\*(.+?)\*\*|__(.+?)__")


def _parse_runs(text: str) -> tuple[Run, ...]:
    """
    Split a template line into styled runs.

    Bold and underline do not nest here: a caption needs "**Cause No: __x__**"
    to come out with the whole line bold and only the number underlined, so a
    bold span is re-scanned for underline spans inside it.
    """
    runs: list[Run] = []

    def emit(chunk: str, bold: bool) -> None:
        if not chunk:
            return
        if bold:
            position = 0
            for match in re.finditer(r"__(.+?)__", chunk):
                if match.start() > position:
                    runs.append(Run(chunk[position:match.start()], bold=True))
                runs.append(Run(match.group(1), bold=True, underline=True))
                position = match.end()
            if position < len(chunk):
                runs.append(Run(chunk[position:], bold=True))
        else:
            runs.append(Run(chunk))

    position = 0
    for match in _MARKUP.finditer(text):
        if match.start() > position:
            emit(text[position:match.start()], False)
        if match.group(1) is not None:
            emit(match.group(1), True)
        else:
            runs.append(Run(match.group(2), underline=True))
        position = match.end()
    if position < len(text):
        emit(text[position:], False)
    return tuple(runs)


def _caption_context(matter: Any, exhibit_name: str) -> tuple[dict[str, str], list[str]]:
    """
    Resolve the template's placeholders against a matter.

    Every value a court document needs and this matter does not have becomes a
    printed blank and a warning. Refusing to build the exhibit would be worse —
    the attorney often wants the numbers long before a cause number exists — and
    printing "None" onto a caption would be worse still.
    """
    warnings: list[str] = []

    def required(value: Optional[str], label: str) -> str:
        if value and value.strip():
            return value.strip()
        warnings.append("No %s on this matter — the exhibit shows a blank." % label)
        return _BLANK

    cause_number = required(getattr(matter, "matter_number", None), "cause number")
    court_name = required(getattr(matter, "court_name", None), "court name")

    case_style = (getattr(matter, "case_style", None) or "").strip()
    if not case_style:
        # The internal short name at least identifies the case, which an empty
        # line does not. Flagged, because it is not caption language.
        case_style = (getattr(matter, "matter_name", None) or "").strip()
        warnings.append(
            "No case style on this matter — using the matter name, which is not "
            "how a caption should read. Set the case style on the matter."
        )

    alignment = getattr(matter, "client_alignment", None)
    if alignment is None:
        # The exhibit is still titled, just not attributed to a side.
        alignment_possessive = ""
        warnings.append(
            "No party alignment on this matter — the exhibit is titled without one. "
            "Set it to title this \"Petitioner's %s\"." % exhibit_name
        )
    else:
        if not isinstance(alignment, ClientAlignment):
            alignment = ClientAlignment(alignment)
        alignment_possessive = "%s's " % alignment.caption

    return {
        "cause_number": cause_number,
        "court_name": court_name,
        "county": (getattr(matter, "county", None) or _BLANK).strip(),
        "state": (getattr(matter, "state", None) or "Texas").strip(),
        "case_style": case_style,
        "alignment_possessive": alignment_possessive,
        "exhibit_name": exhibit_name,
    }, warnings


def caption_lines(
    matter: Any,
    exhibit_name: str,
    template: tuple[tuple[str, bool], ...] = _SYSTEM_CAPTION,
) -> tuple[tuple[Line, ...], list[str]]:
    """
    Build the caption for a matter.

    :param template: The caption to use. Defaults to the firm-wide one; the
        parameter is what a future per-firm override will pass instead.
    :return: ``(lines, warnings)`` — warnings name what the matter could not
        supply, so the UI can say so before anybody prints it.
    :rtype: tuple[tuple[Line, ...], list[str]]
    """
    context, warnings = _caption_context(matter, exhibit_name)
    lines: list[Line] = []
    for text, blank_after in template:
        # Parse the markup FIRST, then substitute into each run. Interpolating
        # before parsing would let a value be read as markup: the blank rule is
        # a row of underscores and would be eaten as an __underline__ marker,
        # and a case style containing ** would corrupt the rest of the caption.
        # A value is content; only the template carries style.
        try:
            runs = tuple(
                Run(run.text.format(**context), bold=run.bold, underline=run.underline)
                for run in _parse_runs(text)
            )
        except KeyError as e:
            LOGGER.error("exhibit_service: caption template references unknown field %s", str(e))
            continue
        if not any(run.text.strip() for run in runs):
            continue
        lines.append(Line(runs=runs, blank_after=blank_after))
    return tuple(lines), warnings


def _plain(line: Line) -> str:
    return "".join(run.text for run in line.runs)


def money(value: Any) -> str:
    """
    Format an amount as currency without ever parsing it as a float.

    The value arrives as a string precisely so exact cents survive Postgres
    ``numeric``; running it through ``float`` to add thousands separators would
    undo that at the last step, in the one place where the figure is about to be
    read into evidence. Grouping is done on the integer part as digits, so the
    fractional part is never arithmetic at all.

    A negative reads ``-$1,200.00`` rather than the accounting parenthesis: a
    minus sign needs no convention explained to whoever is reading the exhibit.
    """
    text = str(value if value is not None else "").strip()
    if not text:
        return ""
    negative = text.startswith("-")
    text = text.lstrip("+-")
    whole, _, fraction = text.partition(".")
    if not whole.isdigit():
        # Not a number after all — hand it back untouched rather than mangling it.
        return str(value)
    grouped = "{:,}".format(int(whole))
    cents = (fraction + "00")[:2]
    return "%s$%s.%s" % ("-" if negative else "", grouped, cents)


def _cell(value: str, column: Column) -> str:
    """One table cell, formatted for an exhibit rather than for a spreadsheet."""
    return money(value) if column.money else (value or "")


# ── Renderers ────────────────────────────────────────────────────────────────

def to_csv(exhibit: Exhibit) -> bytes:
    """
    The clean extraction: header row, data rows, nothing else.

    No caption and no notice, on purpose. This file is meant to be opened in a
    spreadsheet or handed to a model, and a preamble above the header turns a
    valid CSV into something every reader has to be told how to skip. The
    exhibit formats are where the caption and the verification notice live.

    Written with a UTF-8 BOM: Excel reads a plain UTF-8 CSV as the system
    codepage and mangles anything non-ASCII in a payee name.
    """
    buffer = io.StringIO(newline="")
    writer = csv.writer(buffer, lineterminator="\r\n")
    writer.writerow([column.heading for column in exhibit.columns])
    writer.writerows(exhibit.rows)
    return buffer.getvalue().encode("utf-8-sig")


def to_markdown(exhibit: Exhibit) -> bytes:
    """
    The exhibit as markdown — the format meant to be pasted into a model.

    Centring is not expressed. Markdown has no way to say it that survives being
    read as text, and this rendering is optimised for a reader that will lay the
    document out itself: what it needs is the caption's content and the criteria
    that produced the table, not its typography.
    """
    out: list[str] = []
    for line in exhibit.caption:
        out.append(_md_line(line))
        out.append("")

    # The table leads. Selection follows it rather than preceding it: the
    # exhibit's substance is the evidence, and how it was selected is the
    # methodology note a reader turns to afterwards.
    if exhibit.columns:
        out.append("| " + " | ".join(c.heading for c in exhibit.columns) + " |")
        out.append("| " + " | ".join("---:" if c.numeric else "---" for c in exhibit.columns) + " |")
        for row in exhibit.rows:
            out.append("| " + " | ".join(
                _md_cell(_cell(value, column))
                for value, column in zip(row, exhibit.columns)
            ) + " |")
        out.append("")

    if exhibit.footnotes:
        for note in exhibit.footnotes:
            out.append("*%s*" % note)
            out.append("")

    if exhibit.summary:
        out.append("## Totals")
        out.append("")
        for label, value in exhibit.summary:
            out.append("- **%s:** %s" % (label, value))
        out.append("")

    if exhibit.selection:
        out.append("## Selection")
        out.append("")
        for label, value in exhibit.selection:
            out.append("- **%s:** %s" % (label, value))
        out.append("")

    out.append("---")
    out.append("")
    out.append("*%s*" % exhibit.notice)
    out.append("")
    return "\n".join(out).encode("utf-8")


def _md_line(line: Line) -> str:
    """
    One caption line as markdown.

    Adjacent runs sharing a style are merged before the markers go on. Emitting
    per run would close and reopen emphasis mid-phrase — ``**Cause No: ****DF-24
    -01234**`` — which renders as literal asterisks. Underline has no markdown
    spelling and is dropped; bold already carries the emphasis, and the point of
    this rendering is the content, not the typography.
    """
    parts: list[str] = []
    buffer: list[str] = []
    bold: Optional[bool] = None

    def flush() -> None:
        if buffer:
            joined = "".join(buffer)
            parts.append("**%s**" % joined if bold else joined)

    for run in line.runs:
        if bold is not None and run.bold != bold:
            flush()
            buffer.clear()
        bold = run.bold
        buffer.append(run.text)
    flush()
    return "".join(parts)


def _md_cell(value: str) -> str:
    """A pipe inside a description would end the cell early."""
    return (value or "").replace("|", "\\|").replace("\n", " ")


def _docx_page_number(paragraph: Any) -> None:
    """
    Put a live "Page N of M" in a paragraph.

    python-docx has no page-number API because a page number is not text — it is
    a field Word evaluates at layout time, when it finally knows where the pages
    fell. So the field is written as raw OOXML: begin, the instruction, end.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def field(instruction: str) -> None:
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = instruction
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(end)

    paragraph.add_run("Page ")
    field(" PAGE ")
    paragraph.add_run(" of ")
    field(" NUMPAGES ")


def _docx_repeat_header(table: Any) -> None:
    """
    Mark the first row as a header Word repeats on every page.

    A table that runs to a second page without its headings makes the reader
    count columns to find out which one holds the amount. `w:tblHeader` is the
    flag; python-docx does not surface it.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def to_docx(exhibit: Exhibit) -> bytes:
    """The exhibit as a Word document, caption centred and table ruled."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _docx_page_number(footer)
    for run in footer.runs:
        run.font.size = Pt(9)

    for line in exhibit.caption:
        paragraph = document.add_paragraph()
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if line.align == "center" else WD_ALIGN_PARAGRAPH.LEFT
        )
        for run in line.runs:
            drawn = paragraph.add_run(run.text)
            drawn.bold = run.bold
            drawn.underline = run.underline
        if line.blank_after:
            document.add_paragraph()

    if exhibit.columns:
        table = document.add_table(rows=1, cols=len(exhibit.columns))
        table.style = "Table Grid"
        for cell, column in zip(table.rows[0].cells, exhibit.columns):
            cell.text = ""
            run = cell.paragraphs[0].add_run(column.heading)
            run.bold = True
        _docx_repeat_header(table)
        for row in exhibit.rows:
            cells = table.add_row().cells
            for cell, value, column in zip(cells, row, exhibit.columns):
                cell.text = ""
                paragraph = cell.paragraphs[0]
                paragraph.add_run(_cell(value, column))
                if column.numeric:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        document.add_paragraph()

    for note in exhibit.footnotes:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(note)
        run.italic = True
        run.font.size = Pt(8)

    if exhibit.summary:
        heading = document.add_paragraph()
        heading.add_run("Totals").bold = True
        for label, value in exhibit.summary:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run("%s: " % label).bold = True
            paragraph.add_run(value)
        document.add_paragraph()

    if exhibit.selection:
        heading = document.add_paragraph()
        heading.add_run("Selection").bold = True
        for label, value in exhibit.selection:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run("%s: " % label).bold = True
            paragraph.add_run(value)
        document.add_paragraph()

    notice = document.add_paragraph()
    notice_run = notice.add_run(exhibit.notice)
    notice_run.italic = True
    notice_run.font.size = Pt(8)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


_PDF_STYLE = ("<style>"
              "body{font-family:serif;font-size:10pt}"
              "table{width:100%;border-collapse:collapse}"
              "th,td{border:1px solid #666;padding:3px;font-size:8.5pt;text-align:left}"
              "th{background:#eee;font-weight:bold}"
              "td.n,th.n{text-align:right}"
              "p{margin:2pt 0}"
              ".c{text-align:center}"
              ".notice{font-size:7.5pt;font-style:italic;color:#333}"
              "</style>")


def _esc(value: str) -> str:
    return html.escape(value or "", quote=False)


def _pdf_caption(exhibit: Exhibit) -> str:
    parts: list[str] = []
    for line in exhibit.caption:
        inner = "".join(
            "%s%s%s%s%s" % (
                "<b>" if run.bold else "", "<u>" if run.underline else "",
                _esc(run.text),
                "</u>" if run.underline else "", "</b>" if run.bold else "",
            )
            for run in line.runs
        )
        parts.append('<p class="c">%s</p>' % inner)
        if line.blank_after:
            parts.append("<p>&#160;</p>")
    return "".join(parts)


def _pdf_list(title: str, entries: tuple[tuple[str, str], ...]) -> str:
    if not entries:
        return ""
    parts = ["<p>&#160;</p><p><b>%s</b></p>" % _esc(title)]
    parts += ["<p>&#8226; <b>%s:</b> %s</p>" % (_esc(label), _esc(value))
              for label, value in entries]
    return "".join(parts)


def _pdf_header_row(exhibit: Exhibit) -> str:
    return "<tr>%s</tr>" % "".join(
        '<th class="n">%s</th>' % _esc(c.heading) if c.numeric else "<th>%s</th>" % _esc(c.heading)
        for c in exhibit.columns
    )


def _pdf_rows(exhibit: Exhibit) -> list[str]:
    rows = []
    for row in exhibit.rows:
        cells = "".join(
            '<td class="n">%s</td>' % _esc(_cell(value, column)) if column.numeric
            else "<td>%s</td>" % _esc(_cell(value, column))
            for value, column in zip(row, exhibit.columns)
        )
        rows.append("<tr>%s</tr>" % cells)
    return rows


def to_pdf(exhibit: Exhibit) -> bytes:
    """
    The exhibit as a PDF, laid out by PyMuPDF's Story engine.

    PyMuPDF is already a dependency — it is what reads the statements in the
    first place — so this adds no new one. WeasyPrint would render richer CSS
    but needs native Pango and Cairo in the image, a real cost for a caption and
    a ruled table.

    **Story does not repeat a table header across pages.** ``<thead>`` is a
    paged-media idea and the engine has no page model — handed 120 rows it emits
    four pages and pages two through four begin mid-data, leaving the reader to
    count columns to find the amount. So the table is not one story: rows are
    measured a page at a time and each page gets its own small table carrying
    its own header. ``place()`` reports whether the content fitted, which is the
    only measurement needed — try a batch, shrink until it fits, then grow while
    it still does. The fit is exact, not conservative: the 120-row case still
    lands in four pages.
    """
    import pymupdf

    page_rect = pymupdf.paper_rect("letter")
    # One-inch margins, with the bottom pulled up to leave the footer its band.
    content = page_rect + (72, 72, -72, -90)
    header = _pdf_header_row(exhibit)
    rows = _pdf_rows(exhibit)

    buffer = io.BytesIO()
    writer = pymupdf.DocumentWriter(buffer)
    state: dict[str, Any] = {"device": None, "top": content.y0, "pages": 0}

    def begin() -> None:
        state["device"] = writer.begin_page(page_rect)
        state["top"] = content.y0
        state["pages"] += 1

    def remaining() -> Any:
        return pymupdf.Rect(content.x0, state["top"], content.x1, content.y1)

    def flow(markup: str) -> None:
        """Place free-flowing content, continuing onto new pages as needed."""
        if not markup:
            return
        story = pymupdf.Story(html=_PDF_STYLE + markup)
        while True:
            if remaining().height < 24:
                writer.end_page()
                begin()
            more, filled = story.place(remaining())
            story.draw(state["device"])
            state["top"] = bottom(filled) + 6
            if not more:
                return
            writer.end_page()
            begin()

    def bottom(filled: Any) -> float:
        """place() reports the filled area as a plain (x0, y0, x1, y1) tuple."""
        return float(filled[3])

    def fits(batch: list[str], rect: Any):
        story = pymupdf.Story(html=_PDF_STYLE + "<table>" + header + "".join(batch) + "</table>")
        more, filled = story.place(rect)
        return (not more), story, filled

    begin()
    flow(_pdf_caption(exhibit))

    pending = list(rows)
    guess = 40
    while pending:
        # A table needs room for its header plus a row before it is worth
        # starting; otherwise take the next page.
        if remaining().height < 60:
            writer.end_page()
            begin()

        count = min(len(pending), guess)
        ok, story, filled = fits(pending[:count], remaining())
        while not ok and count > 1:
            count = max(1, int(count * 0.85))
            ok, story, filled = fits(pending[:count], remaining())
        while count < len(pending):
            grown_ok, grown, grown_filled = fits(pending[:count + 1], remaining())
            if not grown_ok:
                break
            count += 1
            story, filled = grown, grown_filled

        if not ok and count == 1:
            # One row will not fit even on an empty page — a description long
            # enough to overflow. Draw it anyway rather than looping forever.
            LOGGER.warning("exhibit_service.to_pdf: a single row did not fit a page")

        story.place(remaining())
        story.draw(state["device"])
        state["top"] = bottom(filled) + 6
        pending = pending[count:]
        guess = max(5, count)

    for note in exhibit.footnotes:
        flow('<p class="notice">%s</p>' % _esc(note))
    flow(_pdf_list("Totals", exhibit.summary))
    flow(_pdf_list("Selection", exhibit.selection))
    flow('<p>&#160;</p><p class="notice">%s</p>' % _esc(exhibit.notice))
    writer.end_page()
    writer.close()

    return _stamp_page_numbers(buffer.getvalue())


def _stamp_page_numbers(pdf_bytes: bytes) -> bytes:
    """
    Write "Page N of M" centred in the footer band of every page.

    Done after layout because M is not known until then — the same reason Word
    stores a NUMPAGES field rather than a number.
    """
    import pymupdf

    document = pymupdf.open(stream=pdf_bytes, filetype="pdf")
    try:
        total = document.page_count
        for index, page in enumerate(document, start=1):
            label = "Page %d of %d" % (index, total)
            width = pymupdf.get_text_length(label, fontname="helv", fontsize=9)
            page.insert_text(
                ((page.rect.width - width) / 2, page.rect.height - 50),
                label, fontname="helv", fontsize=9, color=(0.2, 0.2, 0.2),
            )
        return document.tobytes()
    finally:
        document.close()


RENDERERS = {
    "csv": (to_csv, "text/csv", "csv"),
    "md": (to_markdown, "text/markdown; charset=utf-8", "md"),
    "docx": (to_docx,
             "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"),
    "pdf": (to_pdf, "application/pdf", "pdf"),
}


def render(exhibit: Exhibit, fmt: str) -> tuple[bytes, str, str]:
    """
    Render an exhibit.

    :return: ``(content, media_type, filename)``
    :rtype: tuple[bytes, str, str]
    """
    try:
        renderer, media_type, extension = RENDERERS[fmt]
    except KeyError:
        raise ValueError("Unknown export format: %s" % fmt) from None
    content = renderer(exhibit)
    LOGGER.info("exhibit_service.render: format=%s rows=%d bytes=%d",
                fmt, len(exhibit.rows), len(content))
    return content, media_type, "%s.%s" % (exhibit.filename_stem, extension)
